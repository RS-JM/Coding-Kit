"""
DOCX-Parser

Liest ein bestehendes Kandidatenprofil im DOCX-Format.
Extrahiert:
  - Rohen Text (für Claude API)
  - Dokumentstruktur (Überschriften, Abschnitte, Tabellen)
  - Styles/Formatting-Metadaten (Fonts, Farben, Layout)

HINWEIS: Die genauen Sektions-Namen und Felder werden nach der
Template-Analyse der hochgeladenen Profile verfeinert.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json


class DocxParser:
    def __init__(self, pfad: str | Path):
        self.pfad = Path(pfad)
        self.doc = Document(str(self.pfad))
        self._struktur: dict = {}

    # ------------------------------------------------------------------
    # Öffentliche Methoden
    # ------------------------------------------------------------------

    def extrahiere_text(self) -> str:
        """Gibt den vollständigen Text des Dokuments zurück (für Claude API)."""
        zeilen = []
        for absatz in self.doc.paragraphs:
            if absatz.text.strip():
                zeilen.append(absatz.text.strip())
        for tabelle in self.doc.tables:
            for zeile in tabelle.rows:
                row_text = " | ".join(
                    zelle.text.strip() for zelle in zeile.cells if zelle.text.strip()
                )
                if row_text:
                    zeilen.append(row_text)
        return "\n".join(zeilen)

    def extrahiere_struktur(self) -> dict:
        """
        Analysiert die Dokumentstruktur (Überschriften → Abschnitte).
        Gibt ein Dict zurück: { "abschnitt_name": ["Inhalt", ...] }
        """
        struktur = {}
        aktueller_abschnitt = "_header"
        struktur[aktueller_abschnitt] = []

        for absatz in self.doc.paragraphs:
            text = absatz.text.strip()
            if not text:
                continue

            stil = absatz.style.name if absatz.style else ""

            if stil.startswith("Heading") or self._ist_ueberschrift(absatz):
                aktueller_abschnitt = text
                struktur[aktueller_abschnitt] = []
            else:
                struktur[aktueller_abschnitt].append(text)

        self._struktur = struktur
        return struktur

    def extrahiere_style_info(self) -> dict:
        """
        Extrahiert Formatierungs-Metadaten für spätere Template-Nutzung.
        Wichtig um das exakte Layout zu reproduzieren.
        """
        style_info = {
            "seitenraender": self._get_seitenraender(),
            "schriftarten": self._get_schriftarten(),
            "farben": self._get_farben(),
            "absatz_styles": self._get_absatz_styles(),
            "hat_tabellen": len(self.doc.tables) > 0,
            "anzahl_tabellen": len(self.doc.tables),
            "anzahl_abschnitte": len(self.doc.sections),
        }
        return style_info

    def als_json(self) -> str:
        """Gibt Struktur + Style-Info als JSON zurück (für Debugging)."""
        return json.dumps(
            {
                "struktur": self.extrahiere_struktur(),
                "styles": self.extrahiere_style_info(),
            },
            ensure_ascii=False,
            indent=2,
        )

    # ------------------------------------------------------------------
    # Private Hilfsmethoden
    # ------------------------------------------------------------------

    def _ist_ueberschrift(self, absatz) -> bool:
        """Heuristik: Erkennt Überschriften auch ohne Heading-Style."""
        if not absatz.runs:
            return False
        erster_run = absatz.runs[0]
        ist_fett = erster_run.bold
        ist_gross = (
            erster_run.font.size and erster_run.font.size >= Pt(12)
        )
        ist_kurz = len(absatz.text.strip()) < 60
        return bool(ist_fett and ist_kurz) or bool(ist_gross and ist_kurz and ist_fett)

    def _get_seitenraender(self) -> dict:
        sektion = self.doc.sections[0]
        return {
            "oben_cm": round(sektion.top_margin.cm, 2) if sektion.top_margin else None,
            "unten_cm": round(sektion.bottom_margin.cm, 2) if sektion.bottom_margin else None,
            "links_cm": round(sektion.left_margin.cm, 2) if sektion.left_margin else None,
            "rechts_cm": round(sektion.right_margin.cm, 2) if sektion.right_margin else None,
        }

    def _get_schriftarten(self) -> list[str]:
        fonts = set()
        for absatz in self.doc.paragraphs:
            for run in absatz.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        return sorted(fonts)

    def _get_farben(self) -> list[str]:
        farben = set()
        for absatz in self.doc.paragraphs:
            for run in absatz.runs:
                if run.font.color and run.font.color.type is not None:
                    try:
                        rgb = run.font.color.rgb
                        farben.add(str(rgb))
                    except Exception:
                        pass
        return sorted(farben)

    def _get_absatz_styles(self) -> list[dict]:
        styles = []
        gesehen = set()
        for absatz in self.doc.paragraphs:
            stil_name = absatz.style.name if absatz.style else "Normal"
            if stil_name not in gesehen:
                gesehen.add(stil_name)
                styles.append({"name": stil_name})
        return styles
